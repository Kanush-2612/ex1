from docx import Document

# Create a new Word document
doc = Document()

# Title
doc.add_heading("Chest X-Ray Pneumonia Detection Using CNN", level=1)

# Team Members section
doc.add_heading("Team Members:", level=2)
doc.add_paragraph("[Add your team details]")

# Project Statement
doc.add_heading("Project Statement:", level=2)
doc.add_paragraph("The aim of this project is to build a Convolutional Neural Network (CNN) model that "
                  "can accurately classify chest X-ray images to detect pneumonia. This tool aims to support "
                  "healthcare providers by providing a quick and reliable way to identify pneumonia cases based "
                  "on X-ray images.")

# Objectives
doc.add_heading("Objectives:", level=2)
doc.add_paragraph("- To develop a CNN model capable of analyzing X-ray images and determining the presence of pneumonia.")
doc.add_paragraph("- To enhance diagnostic accuracy and reduce the time required for pneumonia detection in clinical settings.")

# Why CNN for Medical Imaging
doc.add_heading("Why CNN for Medical Imaging?", level=2)
doc.add_paragraph("CNNs are highly effective for image classification tasks due to their ability to capture "
                  "spatial hierarchies in images. For medical imaging, CNNs provide accurate results by learning "
                  "important features from X-ray images, which assists in reliable diagnoses.")

# Prerequisites
doc.add_heading("Prerequisites", level=2)

# Tools
doc.add_heading("Tools:", level=3)
doc.add_paragraph("- Jupyter Notebook / Google Colab")
doc.add_paragraph("- Visual Studio Code")

# Technologies
doc.add_heading("Technologies:", level=3)
doc.add_paragraph("- Machine Learning")
doc.add_paragraph("- Deep Learning with CNN")
doc.add_paragraph("- TensorFlow and Keras")

# Libraries
doc.add_heading("Libraries:", level=3)
doc.add_paragraph("- NumPy")
doc.add_paragraph("- Pandas")
doc.add_paragraph("- OpenCV")
doc.add_paragraph("- Matplotlib")
doc.add_paragraph("- Seaborn")
doc.add_paragraph("- scikit-learn")

# Methodology
doc.add_heading("Methodology", level=2)

# Step 1
doc.add_heading("Step 1: Import Libraries", level=3)
doc.add_paragraph("```python\nimport numpy as np\nimport pandas as pd\nimport matplotlib.pyplot as plt\n"
                  "import seaborn as sns\nimport tensorflow as tf\nfrom tensorflow.keras import layers, models\n"
                  "from sklearn.model_selection import train_test_split\n```")

# Step 2
doc.add_heading("Step 2: Load and Preprocess the Dataset", level=3)
doc.add_paragraph("- **Data Collection:** The dataset includes labeled chest X-ray images divided into categories "
                  "for pneumonia and normal cases.\n"
                  "- **Data Augmentation:** Techniques like rotation, zoom, and horizontal flip are applied to enhance "
                  "the model’s robustness.\n"
                  "- **Data Scaling:** Pixel values are scaled to normalize the input for the CNN.")

# Step 3
doc.add_heading("Step 3: Build the CNN Model", level=3)
doc.add_paragraph("- **Layer Structure:** The CNN model consists of convolutional, max-pooling, and dense layers.\n"
                  "- **Activation Function:** ReLU for convolutional layers and softmax for the output layer.\n"
                  "- **Compilation:** The model uses categorical cross-entropy as the loss function and accuracy as the "
                  "evaluation metric.")

# Step 4
doc.add_heading("Step 4: Train the Model", level=3)
doc.add_paragraph("- **Epochs and Batch Size:** Training involves [specify number] epochs and a batch size of "
                  "[specify batch size].\n"
                  "- **Validation:** A portion of the dataset is set aside for validation to monitor performance during training.")

# Step 5
doc.add_heading("Step 5: Model Evaluation", level=3)
doc.add_paragraph("- **Metrics Used:** Accuracy, Precision, Recall, and F1-score are calculated to evaluate the model's "
                  "performance on the test set.\n"
                  "- **Confusion Matrix:** The confusion matrix is plotted to visualize true positives, false positives, true negatives, "
                  "and false negatives.")

# Results
doc.add_heading("Results", level=2)
doc.add_paragraph("The model achieves [state accuracy] accuracy on the test dataset. Additionally, precision and recall "
                  "values indicate the model's effectiveness in identifying pneumonia cases accurately.")

# Conclusion
doc.add_heading("Conclusion", level=2)
doc.add_paragraph("Implementing CNN for chest X-ray analysis provides a valuable tool for pneumonia detection, supporting "
                  "healthcare professionals in making timely and accurate diagnoses.")

# Save the document
file_path = "/mnt/data/Chest_XRay_Pneumonia_Detection_Report.docx"
doc.save(file_path)

file_path
